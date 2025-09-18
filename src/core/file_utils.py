"""
File utility module for the Forensic Psych Report Drafter.
Provides functions for file operations, including reading, writing, and previewing file content.
"""

import logging
import os
from pathlib import Path
from datetime import datetime


def read_file_content(file_path):
    """
    Read the full content of a file.
    
    Args:
        file_path (str): Path to the file to read
        
    Returns:
        str: File content
        
    Raises:
        FileNotFoundError: If the file is not found
        IOError: If there's an issue reading the file
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except FileNotFoundError:
        logging.error(f"File not found: {file_path}")
        raise
    except IOError as e:
        logging.error(f"Error reading file {file_path}: {str(e)}")
        raise


def extract_text_from_pdf(pdf_path, max_pages=None):
    """
    Extract text from a PDF file.
    
    Args:
        pdf_path (str): Path to the PDF file
        max_pages (int, optional): Maximum number of pages to extract. If None, all pages are extracted.
        
    Returns:
        str: Extracted text from the PDF
        
    Raises:
        FileNotFoundError: If the file is not found
        Exception: If there's an issue extracting text from the PDF
    """
    try:
        import fitz  # PyMuPDF
        
        doc = fitz.open(pdf_path)
        num_pages = len(doc)
        
        if max_pages is not None:
            num_pages = min(num_pages, max_pages)
        
        text = []
        for i in range(num_pages):
            page = doc[i]
            text.append(f"--- Page {i+1} ---\n{page.get_text()}")
        
        doc.close()
        
        return "\n\n".join(text)
    except FileNotFoundError:
        logging.error(f"PDF file not found: {pdf_path}")
        raise
    except Exception as e:
        logging.error(f"Error extracting text from PDF {pdf_path}: {str(e)}")
        raise


def read_file_preview(file_path, max_chars=5000, max_lines=None, max_pages=2):
    """
    Read a preview of a file.
    
    Args:
        file_path (str): Path to the file to read
        max_chars (int): Maximum number of characters to read (for text files)
        max_lines (int, optional): Maximum number of lines to read (for text files)
        max_pages (int, optional): Maximum number of pages to extract (for PDF files)
        
    Returns:
        str or tuple: For text files, returns (preview_content, is_truncated).
                      For PDF files, returns extracted text as string.
            
    Raises:
        FileNotFoundError: If the file is not found
        IOError: If there's an issue reading the file
    """
    # Check if the file is a PDF
    if file_path.lower().endswith('.pdf'):
        return extract_text_from_pdf(file_path, max_pages)
    
    # Handle text files
    try:
        if max_lines is not None:
            # Read by lines if max_lines is specified
            with open(file_path, 'r', encoding='utf-8') as file:
                lines = []
                for i, line in enumerate(file):
                    if i >= max_lines:
                        return '\n'.join(lines), True
                    lines.append(line.rstrip('\n'))
                return '\n'.join(lines), False
        else:
            # Read by characters if max_chars is specified
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read(max_chars + 1)  # Read one extra char to check if truncated
                
            if len(content) > max_chars:
                return content[:max_chars], True
            return content, False
    except FileNotFoundError:
        logging.error(f"File not found: {file_path}")
        raise
    except IOError as e:
        logging.error(f"Error reading file {file_path}: {str(e)}")
        raise


def write_file_content(file_path, content):
    """
    Write content to a file.
    
    Args:
        file_path (str): Path to the file to write
        content (str): Content to write to the file
        
    Raises:
        IOError: If there's an issue writing to the file
    """
    try:
        # Ensure the directory exists
        directory = os.path.dirname(file_path)
        if directory and not os.path.exists(directory):
            os.makedirs(directory)
            
        with open(file_path, 'w', encoding='utf-8') as file:
            file.write(content)
    except IOError as e:
        logging.error(f"Error writing to file {file_path}: {str(e)}")
        raise


def suggest_output_filename(input_path, suffix="output", default_name="output.md"):
    """
    Suggest an output filename based on an input filename.
    
    Args:
        input_path (str): Path to the input file
        suffix (str): Suffix to add to the filename
        default_name (str): Default name to use if input_path is None or invalid
        
    Returns:
        str: Suggested output filename
    """
    if not input_path:
        return default_name
        
    try:
        # Get the base filename without extension
        base_name = os.path.basename(input_path)
        name, ext = os.path.splitext(base_name)
        
        # Keep the original extension if it exists, otherwise use .md
        if not ext:
            ext = ".md"
            
        # Create the new filename with the suffix
        return f"{name}-{suffix}{ext}"
    except:
        return default_name


def process_docx_to_markdown(docx_path: str, output_dir: str) -> str:
    """
    Convert a Word document to markdown format.
    
    Args:
        docx_path: Path to the Word document
        output_dir: Directory to save the output
        
    Returns:
        str: Path to the generated markdown file
    """
    try:
        import pypandoc
        
        # Prepare output path
        docx_path = Path(docx_path)
        output_dir = Path(output_dir)
        output_dir.mkdir(exist_ok=True)
        
        output_filename = docx_path.stem + ".md"
        output_path = output_dir / output_filename
        
        # Convert using pandoc
        pypandoc.convert_file(
            str(docx_path),
            'md',
            outputfile=str(output_path),
            extra_args=['--wrap=none', '--extract-media=' + str(output_dir / 'media')]
        )
        
        logging.info(f"Converted {docx_path.name} to markdown")
        return str(output_path)
        
    except ImportError:
        # Fallback: use python-docx to extract text
        logging.warning("pypandoc not available, using basic text extraction")
        try:
            from docx import Document

            docx_path = Path(docx_path)
            output_dir = Path(output_dir)
            output_dir.mkdir(exist_ok=True)

            doc = Document(str(docx_path))
            
            # Extract text from paragraphs
            content = []
            for para in doc.paragraphs:
                if para.text.strip():
                    content.append(para.text)
                    
            # Extract text from tables
            for table in doc.tables:
                content.append("\n| " + " | ".join(["Column"] * len(table.columns)) + " |")
                content.append("| " + " | ".join(["---"] * len(table.columns)) + " |")
                for row in table.rows:
                    row_text = "| " + " | ".join(cell.text.strip() for cell in row.cells) + " |"
                    content.append(row_text)
                content.append("")
                
            # Save as markdown
            output_filename = docx_path.stem + ".md"
            output_path = output_dir / output_filename
            
            markdown_content = "\n\n".join(content)
            
            # Add metadata
            metadata = f"""---
title: {docx_path.stem}
source: {docx_path.name}
converted: {datetime.now().isoformat()}
---

"""
            
            write_file_content(str(output_path), metadata + markdown_content)
            logging.info(f"Extracted text from {docx_path.name} to markdown")
            return str(output_path)
            
        except ImportError:
            raise ImportError("Neither pypandoc nor python-docx is available for Word document processing")
    except Exception as e:
        logging.error(f"Error processing Word document: {e}")
        raise


def process_txt_to_markdown(txt_path: str, output_dir: str) -> str:
    """
    Convert a text file to markdown format.
    
    Args:
        txt_path: Path to the text file
        output_dir: Directory to save the output
        
    Returns:
        str: Path to the generated markdown file
    """
    try:
        txt_path = Path(txt_path)
        output_dir = Path(output_dir)
        output_dir.mkdir(exist_ok=True)
        
        # Read the text file
        content = read_file_content(str(txt_path))
        
        # Add metadata header
        metadata = f"""---
title: {txt_path.stem}
source: {txt_path.name}
converted: {datetime.now().isoformat()}
---

"""
        
        # Basic formatting: treat double newlines as paragraph breaks
        formatted_content = content.replace('\n\n', '\n\n')
        
        # Save as markdown
        output_filename = txt_path.stem + ".md"
        output_path = output_dir / output_filename
        
        write_file_content(str(output_path), metadata + formatted_content)
        logging.info(f"Converted {txt_path.name} to markdown")
        return str(output_path)
        
    except Exception as e:
        logging.error(f"Error processing text file: {e}")
        raise
